"""
redact.py — core PII detection, replacement, and document processing.

This is a deliberately SIMPLE, rule-based implementation for a college
assignment. It uses only regular expressions and small context rules —
no NER model, no external ML library. This is explicit in the design:
a rule-based approach is easy to explain, easy to reproduce, and has
well-understood limitations (see README.md "Limitations").

Sections in this file:
  1. Entity representation + overlap resolution
  2. Structured detectors (EMAIL, PHONE, SSN, CREDIT_CARD, IP, DOB)
  3. Context-rule detectors (PERSON, COMPANY, ADDRESS)
  4. Synthetic replacement (simple dict-based, no external library)
  5. Document processing (DOCX / PDF / TXT -> redacted DOCX)
"""

import os
import re
from dataclasses import dataclass, field

# ---------------------------------------------------------------------
# 1. Entity representation + overlap resolution
# ---------------------------------------------------------------------

@dataclass
class Entity:
    type: str
    text: str
    start: int
    end: int

    def overlaps(self, other):
        return self.start < other.end and other.start < self.end


# Priority used only to decide which entity wins when two different
# types overlap on the same text (higher wins).
TYPE_PRIORITY = {
    "SSN": 100, "CREDIT_CARD": 100, "EMAIL": 95, "IP_ADDRESS": 95,
    "PHONE": 90, "DOB": 85, "ADDRESS": 70, "COMPANY": 60, "PERSON": 50,
}


def merge_entities(entities):
    """Resolve overlapping spans: higher-priority type wins; ties broken
    by longer span. No two entities in the result overlap."""
    seen = {}
    for e in entities:
        key = (e.type, e.start, e.end)
        seen[key] = e
    candidates = list(seen.values())
    candidates.sort(key=lambda e: (-TYPE_PRIORITY.get(e.type, 0), -(e.end - e.start)))
    accepted = []
    for e in candidates:
        if not any(e.overlaps(a) for a in accepted):
            accepted.append(e)
    accepted.sort(key=lambda e: e.start)
    return accepted


# ---------------------------------------------------------------------
# 2. Structured detectors — regex only, no ML
# ---------------------------------------------------------------------

EMAIL_RE = re.compile(r"\b[A-Za-z0-9][A-Za-z0-9._%+\-]*@[A-Za-z0-9.\-]+\.[A-Za-z]{2,}\b")

# Indian mobile (+91/91/0 prefix optional), Indian landline with STD code
# (with or without country code), and a generic US-style fallback.
PHONE_PATTERNS = [
    re.compile(r"(?<!\d)\+?91[\-\s]\d{2,4}[\-\s]\d{3,4}[\-\s]?\d{3,4}(?!\d)"),   # +91 20 4505 3237
    re.compile(r"(?<!\d)0\d{2,4}[\-\s]\d{6,8}(?!\d)"),                          # 022-68052182 (no CC)
    re.compile(r"(?<!\d)(?:\+?91[\-\s]?|0)?[6-9]\d{4}[\-\s]?\d{5}(?!\d)"),      # Indian mobile
    re.compile(r"(?<!\d)(?:\+?1[\-\s]?)?\(?\d{3}\)?[\-\s]\d{3}[\-\s]\d{4}(?!\d)"),  # US-style
]

SSN_RE = re.compile(r"(?<!\d)(\d{3})-(\d{2})-(\d{4})(?!\d)")

CC_CANDIDATE_RE = re.compile(r"(?<!\d)(?:\d[ \-]?){13,19}(?!\d)")

IP_CANDIDATE_RE = re.compile(r"\b(?:\d{1,3}\.){3}\d{1,3}\b")

DOB_CONTEXT_RE = re.compile(r"\b(date of birth|dob|birth ?date|born)\b", re.IGNORECASE)
DATE_RE = re.compile(
    r"\b(\d{1,2}[/\-.]\d{1,2}[/\-.]\d{2,4}"
    r"|\d{1,2}\s+(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Sept|Oct|Nov|Dec)[a-z]*\s+\d{4}"
    r"|(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Sept|Oct|Nov|Dec)[a-z]*\s+\d{1,2},?\s+\d{4})\b",
    re.IGNORECASE,
)


def _luhn_valid(digits: str) -> bool:
    total = 0
    for i, ch in enumerate(digits[::-1]):
        d = int(ch)
        if i % 2 == 1:
            d *= 2
            if d > 9:
                d -= 9
        total += d
    return total % 10 == 0


def detect_structured(text: str):
    out = []
    for m in EMAIL_RE.finditer(text):
        out.append(Entity("EMAIL", m.group(), m.start(), m.end()))

    seen_spans = set()
    for pattern in PHONE_PATTERNS:
        for m in pattern.finditer(text):
            digits = sum(c.isdigit() for c in m.group())
            if not (10 <= digits <= 12):
                continue
            key = (m.start(), m.end())
            if key in seen_spans:
                continue
            seen_spans.add(key)
            out.append(Entity("PHONE", m.group(), m.start(), m.end()))

    for m in SSN_RE.finditer(text):
        area, group, serial = m.group(1), m.group(2), m.group(3)
        if area in ("000", "666") or area.startswith("9") or group == "00" or serial == "0000":
            continue
        out.append(Entity("SSN", m.group(), m.start(), m.end()))

    for m in CC_CANDIDATE_RE.finditer(text):
        digits = re.sub(r"[ \-]", "", m.group())
        if 13 <= len(digits) <= 19 and _luhn_valid(digits):
            out.append(Entity("CREDIT_CARD", m.group(), m.start(), m.end()))

    for m in IP_CANDIDATE_RE.finditer(text):
        parts = m.group().split(".")
        if len(parts) == 4 and all(p.isdigit() and (p == "0" or p[0] != "0") and 0 <= int(p) <= 255 for p in parts):
            out.append(Entity("IP_ADDRESS", m.group(), m.start(), m.end()))

    context_spans = [m.span() for m in DOB_CONTEXT_RE.finditer(text)]
    if context_spans:
        for dm in DATE_RE.finditer(text):
            ds, de = dm.span()
            if any((cs - 40) <= ds and de <= (ce + 40) for cs, ce in context_spans):
                out.append(Entity("DOB", dm.group(), ds, de))

    return out


# ---------------------------------------------------------------------
# 3. Context-rule detectors — PERSON, COMPANY, ADDRESS
#    (basic hybrid logic: regex + surrounding-word context, NOT a
#    general-purpose name model)
# ---------------------------------------------------------------------

# PERSON: only fires next to an explicit title or role label. This is a
# deliberate, narrow rule — it will not catch every name in free-flowing
# prose (see README "False negatives").
PERSON_RE = re.compile(
    r"\b(?:Mr\.|Mrs\.|Ms\.|Dr\.)\s+([A-Z][a-z]+(?:\s+[A-Z][a-z]+){0,3})"
    r"|\b(?:Director|Contact Person|Company Secretary)\s*:\s*([A-Z][a-zA-Z.]+(?:\s+[A-Z][a-zA-Z.]+){0,3})",
)


def detect_person(text: str):
    out = []
    for m in PERSON_RE.finditer(text):
        name = m.group(1) or m.group(2)
        start = m.start(1) if m.group(1) else m.start(2)
        out.append(Entity("PERSON", name, start, start + len(name)))
    return out


COMPANY_SUFFIX = r"(?:Private\s+Limited|Pvt\.?\s*Ltd\.?|Limited|Ltd\.?|LLP|LLC|Inc\.?|Corporation|Corp\.?)"
COMPANY_RE = re.compile(r"\b(?:[A-Z][a-zA-Z&\-']*\s+){1,5}" + COMPANY_SUFFIX + r"\b")
COMPANY_RE_CAPS = re.compile(
    r"\b(?:[A-Z][A-Z&\-']*\s+){1,5}(?:PRIVATE\s+LIMITED|PVT\.?\s*LTD\.?|LIMITED|LTD\.?|LLP|LLC|INC\.?|CORPORATION|CORP\.?)\b"
)
CONNECTOR_RE = re.compile(r"^(?:(?:of|and|the|company)\s+)+", re.IGNORECASE)
SUFFIX_ONLY_WORDS = {"private", "limited", "ltd", "pvt", "llp", "llc", "inc", "corp", "corporation"}


def detect_company(text: str):
    out = []
    seen = set()
    for pattern in (COMPANY_RE, COMPANY_RE_CAPS):
        for m in pattern.finditer(text):
            raw, start = m.group(), m.start()
            trimmed_m = CONNECTOR_RE.match(raw)
            if trimmed_m:
                start += trimmed_m.end()
                raw = raw[trimmed_m.end():]
            words = [w.strip(".") for w in raw.split()]
            if not words or all(w.lower() in SUFFIX_ONLY_WORDS for w in words):
                continue
            end = start + len(raw)
            key = (start, end)
            if key in seen:
                continue
            seen.add(key)
            out.append(Entity("COMPANY", raw.strip(), start, end))
    return out


ADDRESS_ANCHOR_RE = re.compile(
    r"(?:Registered\s+Office|Regd\.?\s+Office|Corporate\s+Office|Registered\s+Address|Address)\s*:?\s*",
    re.IGNORECASE,
)
ADDRESS_KEYWORDS = r"Road|Street|Lane|Nagar|Marg|Plot|Society|Colony|Chowk|Compound"
PINCODE_RE = re.compile(r"\b[1-9]\d{5}\b")
BLOCK_END_RE = re.compile(r"[\n\r]+|(?<=[.;])\s+")
MAX_BLOCK = 220


def detect_address(text: str):
    out = []
    for m in ADDRESS_ANCHOR_RE.finditer(text):
        remainder = text[m.end(): m.end() + MAX_BLOCK]
        bm = BLOCK_END_RE.search(remainder)
        block = remainder[: bm.start()] if bm else remainder
        if not block.strip():
            continue
        if PINCODE_RE.search(block) or re.search(ADDRESS_KEYWORDS, block, re.IGNORECASE):
            start, end = m.end(), m.end() + len(block)
            out.append(Entity("ADDRESS", text[start:end].strip(), start, end))
    return out


def detect_all(text: str):
    entities = detect_structured(text) + detect_company(text) + detect_address(text) + detect_person(text)
    return merge_entities(entities)


# ---------------------------------------------------------------------
# 4. Synthetic replacement — plain dict, no external library
# ---------------------------------------------------------------------

FAKE_FIRST_NAMES = ["Alex", "Jordan", "Sam", "Taylor", "Casey", "Morgan", "Riley", "Avery"]
FAKE_LAST_NAMES = ["Brown", "Smith", "Clark", "Lewis", "Walker", "Young", "King", "Wright"]
FAKE_COMPANY_WORDS = ["Bright", "Summit", "Silver", "Northgate", "Bluewave", "Crestline"]


class Replacer:
    """original_value -> fake_value, consistent within one run.
    The mapping is kept only in memory for the duration of one redaction
    call and is discarded afterward (never written to disk)."""

    def __init__(self):
        self._map = {}
        self._counter = 0

    def _key(self, etype, text):
        if etype == "EMAIL":
            return (etype, text.strip().lower())
        if etype in ("PHONE", "SSN", "CREDIT_CARD"):
            digits = re.sub(r"\D", "", text)
            return (etype, digits[-10:] if etype == "PHONE" else digits)
        return (etype, text.strip())

    def replace(self, entity: Entity) -> str:
        key = self._key(entity.type, entity.text)
        if key in self._map:
            return self._map[key]
        fake = self._generate(entity)
        self._map[key] = fake
        return fake

    def _generate(self, entity: Entity) -> str:
        self._counter += 1
        n = self._counter
        t = entity.type
        if t == "PERSON":
            name = f"{FAKE_FIRST_NAMES[n % len(FAKE_FIRST_NAMES)]} {FAKE_LAST_NAMES[n % len(FAKE_LAST_NAMES)]}"
            return name.upper() if entity.text.isupper() else name
        if t == "EMAIL":
            return f"user{n}@example.com"
        if t == "PHONE":
            is_indian = entity.text.strip().startswith("+91") or entity.text.strip().startswith("91")
            fake_digits = f"{6 + (n % 4)}{100000000 + n * 7 % 89999999:08d}"[:10]
            return f"+91 {fake_digits[:5]} {fake_digits[5:]}" if is_indian else f"({200+n%700}) 555-{1000+n:04d}"
        if t == "COMPANY":
            suffix = "Limited"
            for s in ("Private Limited", "Pvt Ltd", "Limited", "Ltd", "LLP", "LLC", "Inc", "Corp", "Corporation"):
                if s.lower() in entity.text.lower():
                    suffix = s
                    break
            name = f"{FAKE_COMPANY_WORDS[n % len(FAKE_COMPANY_WORDS)]} {suffix}"
            return name.upper() if entity.text.isupper() else name
        if t == "ADDRESS":
            return f"{100+n} Example Street, Example City - {100000+n*7 % 899999}"
        if t == "SSN":
            return f"{100+n%800:03d}-{10+n%80:02d}-{1000+n:04d}"
        if t == "CREDIT_CARD":
            partial = "4" + "".join(str((n * 7 + i) % 10) for i in range(14))
            # Luhn check-digit computation so the fake number is well-formed
            digits = [int(d) for d in partial][::-1]
            s = 0
            for i, d in enumerate(digits):
                if i % 2 == 0:
                    d *= 2
                    if d > 9:
                        d -= 9
                s += d
            check = (10 - (s % 10)) % 10
            full = partial + str(check)
            return " ".join(full[i:i + 4] for i in range(0, len(full), 4))
        if t == "DOB":
            return f"{1 + n % 28:02d}/{1 + n % 12:02d}/{1970 + n % 30}"
        if t == "IP_ADDRESS":
            block = ["192.0.2", "198.51.100", "203.0.113"][n % 3]
            return f"{block}.{1 + n % 250}"
        return "[REDACTED]"


# ---------------------------------------------------------------------
# 5. Document processing — DOCX (structure-preserving), PDF, TXT
# ---------------------------------------------------------------------

SUPPORTED_EXTENSIONS = {".docx", ".pdf", ".txt"}


def _paragraph_run_offsets(paragraph):
    """Map a paragraph's concatenated run text back to individual runs,
    so we can replace PII in place without rebuilding the document."""
    full_text, offsets, pos = [], [], 0
    for run in paragraph.runs:
        t = run.text
        if t == "":
            continue
        offsets.append((run, pos, pos + len(t)))
        full_text.append(t)
        pos += len(t)
    return "".join(full_text), offsets


def _apply_replacement(offsets, start, end, fake_text):
    touched = [(r, rs, re_) for r, rs, re_ in offsets if rs < end and start < re_]
    if not touched:
        return
    if len(touched) == 1:
        run, rs, re_ = touched[0]
        ls, le = start - rs, end - rs
        run.text = run.text[:ls] + fake_text + run.text[le:]
        return
    first_run, first_start, _ = touched[0]
    last_run, last_start, _ = touched[-1]
    prefix = first_run.text[: start - first_start]
    suffix = last_run.text[end - last_start:]
    first_run.text = prefix + fake_text
    for run, _, _ in touched[1:-1]:
        run.text = ""
    if last_run is not first_run:
        last_run.text = suffix


def _process_paragraph(paragraph, replacer, stats):
    stats["paragraphs_scanned"] += 1
    full_text, offsets = _paragraph_run_offsets(paragraph)
    if not full_text.strip() or not offsets:
        return
    entities = detect_all(full_text)
    if not entities:
        return
    for entity in sorted(entities, key=lambda e: e.start, reverse=True):
        fake_text = replacer.replace(entity)
        _apply_replacement(offsets, entity.start, entity.end, fake_text)
        stats["entities_by_type"][entity.type] = stats["entities_by_type"].get(entity.type, 0) + 1
    stats["paragraphs_modified"] += 1


def redact_docx(input_path: str, output_path: str) -> dict:
    from docx import Document

    document = Document(input_path)
    replacer = Replacer()
    stats = {"paragraphs_scanned": 0, "paragraphs_modified": 0, "tables_scanned": 0,
              "headers_scanned": 0, "footers_scanned": 0, "entities_by_type": {}}

    for paragraph in document.paragraphs:
        _process_paragraph(paragraph, replacer, stats)

    seen_cells = set()
    for table in document.tables:
        stats["tables_scanned"] += 1
        for row in table.rows:
            for cell in row.cells:
                if id(cell._tc) in seen_cells:
                    continue  # merged cells: don't reprocess the same cell twice
                seen_cells.add(id(cell._tc))
                for paragraph in cell.paragraphs:
                    _process_paragraph(paragraph, replacer, stats)

    seen_parts = set()
    for section in document.sections:
        for part, key in ((section.header, "headers_scanned"), (section.footer, "footers_scanned")):
            if part is None or id(part._element) in seen_parts:
                continue
            seen_parts.add(id(part._element))
            stats[key] += 1
            for paragraph in part.paragraphs:
                _process_paragraph(paragraph, replacer, stats)

    stats["total_entities_replaced"] = sum(stats["entities_by_type"].values())
    document.save(output_path)
    return stats


def redact_txt(input_path: str, output_path: str) -> dict:
    from docx import Document

    with open(input_path, "r", encoding="utf-8", errors="replace") as f:
        raw_text = f.read()

    replacer = Replacer()
    stats = {"paragraphs_scanned": 0, "paragraphs_modified": 0, "tables_scanned": 0,
              "headers_scanned": 0, "footers_scanned": 0, "entities_by_type": {}}

    document = Document()
    for line in raw_text.splitlines():
        stats["paragraphs_scanned"] += 1
        entities = detect_all(line)
        redacted = line
        for entity in sorted(entities, key=lambda e: e.start, reverse=True):
            fake = replacer.replace(entity)
            redacted = redacted[: entity.start] + fake + redacted[entity.end:]
            stats["entities_by_type"][entity.type] = stats["entities_by_type"].get(entity.type, 0) + 1
        if redacted != line:
            stats["paragraphs_modified"] += 1
        document.add_paragraph(redacted)

    stats["total_entities_replaced"] = sum(stats["entities_by_type"].values())
    document.save(output_path)
    return stats


def redact_pdf(input_path: str, output_path: str) -> dict:
    from docx import Document
    from pypdf import PdfReader

    reader = PdfReader(input_path)
    replacer = Replacer()
    stats = {"paragraphs_scanned": 0, "paragraphs_modified": 0, "tables_scanned": 0,
              "headers_scanned": 0, "footers_scanned": 0, "entities_by_type": {}}

    document = Document()
    for page_num, page in enumerate(reader.pages, start=1):
        document.add_heading(f"Page {page_num}", level=3)
        for line in (page.extract_text() or "").splitlines():
            stats["paragraphs_scanned"] += 1
            entities = detect_all(line)
            redacted = line
            for entity in sorted(entities, key=lambda e: e.start, reverse=True):
                fake = replacer.replace(entity)
                redacted = redacted[: entity.start] + fake + redacted[entity.end:]
                stats["entities_by_type"][entity.type] = stats["entities_by_type"].get(entity.type, 0) + 1
            if redacted != line:
                stats["paragraphs_modified"] += 1
            document.add_paragraph(redacted)

    stats["total_entities_replaced"] = sum(stats["entities_by_type"].values())
    document.save(output_path)
    return stats


def redact_document(input_path: str, output_path: str) -> dict:
    ext = os.path.splitext(input_path)[1].lower()
    if ext == ".docx":
        return redact_docx(input_path, output_path)
    if ext == ".pdf":
        return redact_pdf(input_path, output_path)
    if ext == ".txt":
        return redact_txt(input_path, output_path)
    raise ValueError(f"Unsupported file type: {ext!r}")
