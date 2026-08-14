import json
import os

from redact import detect_all

GOLD_PATH = os.path.join(os.path.dirname(__file__), "evaluation", "gold.json")
PROSPECTUS_PATH = os.path.join(os.path.dirname(__file__), "Red_Herring_Prospectus.docx")

ALL_TYPES = ["PERSON", "EMAIL", "PHONE", "COMPANY", "ADDRESS", "SSN", "CREDIT_CARD", "DOB", "IP_ADDRESS"]


def _load_gold():
    with open(GOLD_PATH, "r", encoding="utf-8") as f:
        return json.load(f)


def _all_text_units(docx_path):
    """Every paragraph/cell-paragraph text in the real document, in the
    order python-docx exposes them — used to find each gold entity's
    real surrounding context."""
    from docx import Document

    doc = Document(docx_path)
    units = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if p.text.strip():
                        units.append(p.text)
    return units


def _find_context(gold_text, units):
    for u in units:
        if gold_text in u:
            return u
    return None


def evaluate():
    gold = _load_gold()
    units = _all_text_units(PROSPECTUS_PATH)

    results = {t: {"tp": 0, "fp": 0, "fn": 0} for t in ALL_TYPES}
    fp_examples, fn_examples = [], []

    for g in gold:
        context = _find_context(g["text"], units)
        if context is None:
            # Gold text not found verbatim anymore (should not happen if
            # gold.json was built correctly) -> count as a miss.
            results[g["type"]]["fn"] += 1
            fn_examples.append(g)
            continue

        start = context.find(g["text"])
        end = start + len(g["text"])
        predictions = detect_all(context)

        matched = any(
            p.type == g["type"] and p.start < end and start < p.end
            for p in predictions
        )
        if matched:
            results[g["type"]]["tp"] += 1
        else:
            results[g["type"]]["fn"] += 1
            fn_examples.append(g)

        # Any prediction of this gold entity's type in this context that
        # does NOT overlap any gold entity of the same type in this
        # context is a false positive.
        gold_spans_same_context_and_type = [
            (gg["text"], context.find(gg["text"]))
            for gg in gold
            if gg["type"] == g["type"] and gg["text"] in context
        ]
        for p in predictions:
            if p.type != g["type"]:
                continue
            overlaps_any_gold = any(
                gs != -1 and p.start < gs + len(gt) and gs < p.end
                for gt, gs in gold_spans_same_context_and_type
            )
            if not overlaps_any_gold:
                fp_examples.append({"type": p.type, "text": p.text, "context_snippet": context[:80]})

    # Deduplicate FP examples (same prediction can surface once per gold
    # entity sharing that context) and count them per type.
    unique_fps = {}
    for fp in fp_examples:
        key = (fp["type"], fp["text"])
        if key not in unique_fps:
            unique_fps[key] = fp
            results[fp["type"]]["fp"] += 1

    return results, fn_examples, list(unique_fps.values())


def _safe_div(a, b):
    return a / b if b else None


def print_report():
    results, fn_examples, fp_examples = evaluate()
    gold = _load_gold()
    gold_counts = {t: sum(1 for g in gold if g["type"] == t) for t in ALL_TYPES}

    print(f"Target gold-set size: 200 entities.")
    print(f"Actual gold-set size: {len(gold)} entities.\n")

    print(f"{'TYPE':12s} {'Gold':>5s} {'TP':>4s} {'FP':>4s} {'FN':>4s} {'Prec':>7s} {'Recall':>7s} {'F1':>7s}")
    total_tp = total_fp = total_fn = 0
    for t in ALL_TYPES:
        r = results[t]
        tp, fp, fn = r["tp"], r["fp"], r["fn"]
        total_tp += tp; total_fp += fp; total_fn += fn
        p = _safe_div(tp, tp + fp)
        rec = _safe_div(tp, tp + fn)
        f1 = _safe_div(2 * p * rec, p + rec) if (p is not None and rec is not None and (p + rec) > 0) else (0.0 if p is not None and rec is not None else None)
        fmt = lambda v: f"{v:.2f}" if v is not None else "N/A"
        print(f"{t:12s} {gold_counts[t]:5d} {tp:4d} {fp:4d} {fn:4d} {fmt(p):>7s} {fmt(rec):>7s} {fmt(f1):>7s}")

    overall_p = _safe_div(total_tp, total_tp + total_fp)
    overall_r = _safe_div(total_tp, total_tp + total_fn)
    overall_f1 = _safe_div(2 * overall_p * overall_r, overall_p + overall_r) if overall_p and overall_r else None
    accuracy = _safe_div(total_tp, total_tp + total_fp + total_fn)

    print(f"\nTOTAL: TP={total_tp} FP={total_fp} FN={total_fn}")
    print(f"Overall Precision: {overall_p*100:.1f}%" if overall_p is not None else "Overall Precision: N/A")
    print(f"Overall Recall: {overall_r*100:.1f}%" if overall_r is not None else "Overall Recall: N/A")
    print(f"Overall F1: {overall_f1*100:.1f}%" if overall_f1 is not None else "Overall F1: N/A")
    print(f"Entity-level Accuracy (TP/(TP+FP+FN)): {accuracy*100:.1f}%" if accuracy is not None else "Accuracy: N/A")

    print("\n--- False Positives ---")
    for fp in fp_examples:
        print(f"  [{fp['type']}] {fp['text']!r}  (in: ...{fp['context_snippet']}...)")

    print("\n--- False Negatives ---")
    for fn in fn_examples:
        print(f"  [{fn['type']}] {fn['text']!r}  ({fn.get('section','')})")


if __name__ == "__main__":
    print_report()
